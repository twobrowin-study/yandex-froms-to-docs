import io
from docx import (
    Document,
    document, table
)
from docx.text import paragraph

def DocxFillTemplate(template_file_path: str, variables: dict) -> bytes:
    """
    Запись в шаблон документа и возвращает его в байтовом формате
    """
    template_document: document.Document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            _replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for _table in template_document.tables:
            for col in _table.columns:
                for cell in col.cells:
                    cell: table._Cell
                    for paragraph in cell.paragraphs:
                        _replace_text_in_paragraph(paragraph, variable_key, variable_value)

    document_buf = io.BytesIO()

    template_document.save(document_buf)
    
    document_buf.seek(0)
    return document_buf.read()

def _replace_text_in_paragraph(paragraph: paragraph.Paragraph, key: str, value: str):
    """
    Наивная замена текста в параграфе
    """
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)